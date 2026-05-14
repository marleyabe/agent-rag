from __future__ import annotations

import sys
from pathlib import Path

import fitz
import pytest
from docx import Document as DocxDocument


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if not SRC.exists():
    # mutmut runs tests inside a "mutants" workspace sibling to the project root.
    SRC = ROOT.parent / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    file_path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text(
        (72, 72),
        "Contrato de servico\nA multa por atraso e de 2%\nPrazo de pagamento: 10 dias",
    )
    page2 = doc.new_page()
    page2.insert_text(
        (72, 72),
        "Pagamento mensal\nEm caso de inadimplencia, ha nova multa.",
    )
    doc.save(file_path)
    doc.close()
    return file_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    file_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("Contrato principal")
    doc.add_paragraph("A multa por atraso e de 2%.")
    doc.add_paragraph("O pagamento deve ocorrer em 10 dias.")
    doc.save(file_path)
    return file_path

from __future__ import annotations

from pathlib import Path

import pytest

from ingestion.loaders import DocumentLoader


def test_extract_pdf_with_minimum_metadata(sample_pdf: Path) -> None:
    loader = DocumentLoader()
    sections = loader.load(sample_pdf)

    assert len(sections) >= 2
    assert sections[0].file_name == "sample.pdf"
    assert sections[0].file_type == "pdf"
    assert sections[0].page_number == 1
    assert sections[0].paragraph_number is None
    assert sections[0].file_id
    assert "multa" in " ".join(section.text.lower() for section in sections)


def test_extract_docx_with_minimum_metadata(sample_docx: Path) -> None:
    loader = DocumentLoader()
    sections = loader.load(sample_docx)

    assert len(sections) == 3
    assert sections[1].file_name == "sample.docx"
    assert sections[1].file_type == "docx"
    assert sections[1].paragraph_number == 2
    assert sections[1].page_number is None
    assert sections[1].file_id


def test_fail_on_unsupported_file_type(tmp_path: Path) -> None:
    unsupported = tmp_path / "sample.txt"
    unsupported.write_text("abc", encoding="utf-8")
    loader = DocumentLoader()

    with pytest.raises(ValueError, match="Unsupported file type"):
        loader.load(unsupported)


def test_skip_empty_pdf_pages_and_docx_paragraphs(tmp_path: Path) -> None:
    import fitz
    from docx import Document as DocxDocument

    pdf = tmp_path / "blank_page.pdf"
    doc = fitz.open()
    doc.new_page()  # em branco
    p2 = doc.new_page()
    p2.insert_text((72, 72), "conteudo")
    doc.save(pdf)
    doc.close()

    d = tmp_path / "blank_par.docx"
    docx = DocxDocument()
    docx.add_paragraph("")
    docx.add_paragraph("conteudo docx")
    docx.save(d)

    loader = DocumentLoader()
    pdf_sections = loader.load(pdf)
    docx_sections = loader.load(d)
    assert len(pdf_sections) == 1 and pdf_sections[0].page_number == 2
    assert len(docx_sections) == 1 and docx_sections[0].paragraph_number == 2
